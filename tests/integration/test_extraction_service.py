"""Test integracyjny ExtractionService (krok 2.3.2) — uderza w realny kontener Tika.

Symetrycznie do `test_tika_client.py` (transport), tu testujemy warstwe DOMENOWA nad
realnym transportem: `ExtractionService(TikaClient(real))`. Cel, ktorego testy jednostkowe
(z atrapa) NIE dowioda: czy zalozenia domeny o KSZTALCIE metadanych Tiki trzymaja sie
zywego wyjscia — przede wszystkim klucz `Content-Type` i to, ze `_pick_content_type`
poprawnie odcina parametry typu (`; charset=...`) na realnej odpowiedzi Tiki.

Gdy Tika niedostepna -> SKIP (fixture `tika_url`), nie fail. `extract` jest async ->
`asyncio.run` (jak `test_tika_client.py`/`test_llm.py`; projekt nie uzywa pytest-asyncio).
"""

import asyncio
import io
import os

import docx
import pytest
from PIL import Image, ImageDraw, ImageFont

from app.extraction import ExtractionResult, ExtractionService, TikaClient

# Parasol `integration` + wezszy `integration_tika` (uderzamy w kontener Tika).
pytestmark = [pytest.mark.integration, pytest.mark.integration_tika]

# Fonty z polskimi glifami (jak w test_fastapi_extract.py) — pod render stron do OCR.
_FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/TTF/DejaVuSans.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
]


def _find_font() -> str | None:
    return next((p for p in _FONT_CANDIDATES if os.path.exists(p)), None)


def test_extract_docx_przez_serwis(tika_url):
    """DOCX -> ExtractionService: tekst znormalizowany (polskie znaki) + policzone metadane."""
    document = docx.Document()
    document.add_paragraph("Zażółć gęślą jaźń")
    buf = io.BytesIO()
    document.save(buf)

    service = ExtractionService(TikaClient(base_url=tika_url))
    # Bez content_type — autodetekcja typu po stronie Tiki (jak w tescie transportu).
    result = asyncio.run(service.extract(data=buf.getvalue()))

    assert isinstance(result, ExtractionResult)
    assert "gęślą" in result.text                              # polskie znaki przeszly przez normalizacje
    assert result.text == result.text.strip()                  # normalizacja zdjela wiodace/koncowe biale znaki
    assert "wordprocessing" in (result.metadata.content_type or "")  # realny MIME DOCX z metadanych Tiki
    assert result.metadata.char_count > 0
    assert result.metadata.word_count == 3
    # Sentinel (zweryfikowane empirycznie 2026-06-21): Tika w naszej konfiguracji NIE
    # auto-wykrywa jezyka — dokument bez zadeklarowanego jezyka -> language None. Gdy
    # ten assert kiedys padnie (np. po wlaczeniu detekcji jezyka), to sygnal, by wrocic
    # do decyzji "jezyk = best-effort/refinement" z 2.3.2, a nie cichy zwrot wartosci.
    assert result.metadata.language is None


def test_extract_txt_odcina_charset_z_mime(tika_url):
    """text/plain -> ExtractionService: realna Tika zwraca MIME z '; charset=...';

    dowodzimy, ze `_pick_content_type` odcina parametry na ZYWYM wyjsciu (nie tylko na
    atrapie) — content_type to czyste 'text/plain', bez charset.
    """
    service = ExtractionService(TikaClient(base_url=tika_url))
    # Podajemy content_type, by Tika potraktowala bajty jako tekst (bez zgadywania typu).
    result = asyncio.run(
        service.extract(
            data="Pismo do dekretacji.\nDruga linia.".encode("utf-8"),
            content_type="text/plain",
        )
    )

    assert result.metadata.content_type == "text/plain"   # parametry typu (charset) odciete
    assert "dekretacji" in result.text
    assert result.metadata.word_count == 5


def test_extract_jezyk_z_wlasciwosci_docx(tika_url):
    """DOCX z ZADEKLAROWANYM jezykiem -> serwis zwraca go w metadanych (sciezka dc:language).

    Dopelnienie sentinela z `test_extract_docx_przez_serwis` (brak jezyka -> None): tu
    dowodzimy END-TO-END na realnej Tice, ze GDY dokument deklaruje jezyk we wlasciwosciach,
    Tika wystawia go jako `dc:language`, a `_pick_language` go podchwytuje (czego atrapa
    nie zweryfikuje — to zalozenie o realnym kluczu metadanej).
    """
    document = docx.Document()
    document.core_properties.language = "pl"   # jezyk w wlasciwosciach dokumentu -> dc:language
    document.add_paragraph("Pismo do dekretacji w sprawie podatku.")
    buf = io.BytesIO()
    document.save(buf)

    service = ExtractionService(TikaClient(base_url=tika_url))
    result = asyncio.run(service.extract(data=buf.getvalue()))

    assert result.metadata.language == "pl"


def test_extract_skan_png_raportuje_ocr_used(tika_url):
    """Skan-OBRAZ (PNG) -> `ocr_used=True` i MIME bez wewnetrznego znacznika Tiki.

    Regresja: dopoki `ocr_used` czytalo WYLACZNIE `pdf:ocrPageCount`, obrazy raportowaly
    `ocr_used=False` mimo realnego OCR (Tika nie wystawia tej metadanej dla nie-PDF; jedynym
    sladem jest TesseractOCRParser w `X-TIKA:Parsed-By`). Dowod wymaga REALNEJ Tiki —
    atrapa potwierdzilaby tylko nasze zalozenie o ksztalcie metadanych, nie samo zalozenie.
    Przy okazji: Tika znakuje MIME zOCR-owanego obrazu jako `image/ocr-png` (typ spoza IANA),
    a serwis ma go znormalizowac do `image/png`.
    """
    font_path = _find_font()
    if not font_path:
        pytest.skip("Brak fontu TTF z polskimi glifami (np. DejaVuSans) — pomijam OCR.")

    img = Image.new("RGB", (1000, 300), "white")
    ImageDraw.Draw(img).text((40, 100), "Wezwanie", fill="black", font=ImageFont.truetype(font_path, 80))
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    service = ExtractionService(TikaClient(base_url=tika_url))
    result = asyncio.run(service.extract(data=buf.getvalue(), content_type="image/png"))

    assert "ezwanie" in result.text.lower()          # rdzen slowa — odporne na pomylki OCR
    assert result.metadata.ocr_used is True          # sedno regresji
    assert result.metadata.content_type == "image/png"   # nie "image/ocr-png"


def test_extract_docx_nie_raportuje_ocr(tika_url):
    """DOCX (ekstrakcja natywna) -> `ocr_used=False` — straznik przed falszywym pozytywem.

    Sygnal obrazowy (`X-TIKA:Parsed-By`) dokladamy sumą logiczną do sygnalu PDF-owego, wiec
    trzeba pilnowac, ze dokument bez OCR nadal raportuje `False`.
    """
    document = docx.Document()
    document.add_paragraph("Pismo czytane natywnie, bez OCR.")
    buf = io.BytesIO()
    document.save(buf)

    service = ExtractionService(TikaClient(base_url=tika_url))
    result = asyncio.run(service.extract(data=buf.getvalue()))

    assert result.metadata.ocr_used is False


def _make_scan_pdf(pages_text: list[str], font_path: str) -> bytes:
    """Zbuduj PDF OBRAZOWY (bez warstwy tekstowej): kazda strona to render slowa do obrazu.

    Taki PDF MUSI przejsc przez OCR (auto go zOCR-uje) — idealny do testu limitu stron:
    OCR jest kosztowny, wiec wlasnie jego dotyczy `MAX_OCR_PAGES`.
    """
    font = ImageFont.truetype(font_path, 80)
    images = []
    for text in pages_text:
        img = Image.new("RGB", (1000, 300), "white")
        ImageDraw.Draw(img).text((40, 100), text, fill="black", font=font)
        images.append(img)
    buf = io.BytesIO()
    # Wielostronicowy PDF z obrazow (save_all + append_images) — strony to czyste skany.
    images[0].save(buf, format="PDF", save_all=True, append_images=images[1:])
    return buf.getvalue()


def test_extract_limit_stron_tnie_skan_pdf(tika_url):
    """Skan PDF wielostronicowy + `max_ocr_pages=1` -> OCR tylko 1. strony; limit w metadanych.

    Dowodzi straznika zasobow (krok 2.3.5, strategia B "limit PRZED auto") na realnej Tice:
    3-stronicowy PDF obrazowy z limitem 1 ma wyekstrahowac tylko tresc strony 1 (bo do Tiki
    poszla tylko ona), a metadane maja NIE CICHO zaraportowac ciecie (1 z 3 stron).
    """
    font_path = _find_font()
    if not font_path:
        pytest.skip("Brak fontu TTF z polskimi glifami (np. DejaVuSans) — pomijam OCR.")

    # Trzy odrebne, latwe do OCR slowa — po jednym na strone.
    pdf = _make_scan_pdf(["Pierwsza", "Druga", "Trzecia"], font_path)

    # Limit 1 strona: do Tiki ma trafic tylko strona 1 (reszta uciecie PRZED ekstrakcja).
    service = ExtractionService(TikaClient(base_url=tika_url), max_ocr_pages=1)
    result = asyncio.run(service.extract(data=pdf, content_type="application/pdf"))

    # Asercje na RDZENIACH slow — odporne na typowe pomylki OCR (np. "Pierwsza"->"Plerwsza").
    low = result.text.lower()
    assert "rwsza" in low                               # strona 1 ("pierwsza") zOCR-owana
    assert "rug" not in low                             # strona 2 ("druga") NIE poszla (ucieta)
    assert "rzeci" not in low                           # strona 3 ("trzecia") j.w.
    # Limit nie cichy: diagnostyka stron w metadanych.
    assert result.metadata.pages_total == 3
    assert result.metadata.pages_processed == 1
    assert result.metadata.ocr_truncated is True
    assert result.metadata.ocr_used is True             # skan -> OCR realnie poszedl
